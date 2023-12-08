from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Cm
from docx.shared import Pt
from docx.enum.section import WD_ORIENT


def generate_report(count_of_row, count_of_col, col_names, data, doc_info):
    # создание документа
    doc = Document()
    # доступ к первой секции:
    section = doc.sections[0]
    # высота листа в сантиметрах
    section.page_height = Cm(25)
    # ширина листа в сантиметрах
    section.page_width = Cm(30.0)
    section.left_margin = Mm(20.4)
    section.right_margin = Mm(10)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(10)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    # Заголовок
    title = doc.add_heading('Отчёт', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
    title = doc.add_heading(f"за период {doc_info['start_data']} - {doc_info['end_data']}", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in title.runs:
        run.font.size = Pt(15)

    empty_space = doc.add_paragraph(" ")

    # создание таблицы
    table = doc.add_table(rows=count_of_row+1, cols=count_of_col)
    table.style = 'Light Shading Accent 1'

    #заполнение названий столбцов
    for i in range(count_of_col):
        cell = table.cell(0, i)
        cell.paragraphs[0].add_run(col_names[i]).bold = True

    sum_of_types = {}
    # заполнение данными
    num_of_row = 1

    for elem in data:
        num_of_col = 0
        for key, value in elem.items():
            cell = table.cell(num_of_row, num_of_col)
            cell.text = str(value)
            num_of_col += 1
        num_of_row += 1

    sum_of_types = {}
    type_op = ''
    sum_op = 0
    for elem in data:
        if elem["type_of_operation"] == type_op:
            sum_op += elem["sum"]
        else:
            if sum_op != 0:
                sum_of_types[type_op] = sum_op
            type_op = elem["type_of_operation"]
            sum_op = 0
    sum_of_types[type_op] = sum_op

    empty_space = doc.add_paragraph(" ")


    for key, value in sum_of_types.items():
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"Итоги {key}: {value}")
        font = run.font
        font.size = Pt(15)

    empty_space = doc.add_paragraph(" ")

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Утверждаю    {doc_info['signature']}")
    font = run.font
    font.size = Pt(15)

    doc.save("report.docx")


names_of_col = ["Продукт", "Тип операции", "Дата", "Склад", "Количество", "Единица измерения", "Цена за единицу", "Сумма"]

data_tuple = (
    {
        "type": "песок",
        "type_of_operation": "добыча",
        "data": "2023-11-01",
        "stock": "карьер №1",
        "count": 500,
        "unit": "тонн",
        "price": 1000,
        "sum": 500000
    },
    {
        "type": "глина",
        "type_of_operation": "добыча",
        "data": "2023-11-03",
        "stock": "карьер №2",
        "count": 300,
        "unit": "тонн",
        "price": 800,
        "sum": 240000
    },
    {
        "type": "песок",
        "type_of_operation": "продажа",
        "data": "2023-11-05",
        "stock": "потребитель №1",
        "count": 400,
        "unit": "тонн",
        "price": 2000,
        "sum": 800000
    },
    {
        "type": "глина",
        "type_of_operation": "продажа",
        "data": "2023-11-07",
        "stock": "потребитель №2",
        "count": 200,
        "unit": "тонн",
        "price": 1500,
        "sum": 300000
    },
    {
        "type": "песок",
        "type_of_operation": "продажа",
        "data": "2023-11-20",
        "stock": "потребитель №3",
        "count": 600,
        "unit": "тонн",
        "price": 2500,
        "sum": 1500000
    }
)

doc_info = {"start_data": '12.12.2021',
            "end_data": '12.12.2022',
            "signature": "Иванов"}

generate_report(len(data_tuple), len(names_of_col), names_of_col, data_tuple, doc_info)

